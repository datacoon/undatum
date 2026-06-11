"""Document extraction command module."""

import glob
import os
import sys
import tempfile
from typing import Any, Optional

import pandas as pd

from ..utils import get_option
from .packager import Packager

SUPPORTED_INPUTS = {"pdf", "doc", "docx", "xls", "xlsx"}
SUPPORTED_OUTPUTS = {"csv", "json", "ndjson", "parquet", "datapackage"}
OUTPUT_EXT = {
    "csv": ".csv",
    "json": ".json",
    "ndjson": ".ndjson",
    "parquet": ".parquet",
}


def _normalize_output_format(value: Optional[str]) -> str:
    if not value:
        return "csv"
    value = value.lower()
    if value == "jsonl":
        return "ndjson"
    return value


def _normalize_method(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    return value.lower()


def _parse_pages(value: Optional[str]) -> Optional[list[int]]:
    if not value:
        return None
    pages: set[int] = set()
    for part in value.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_str, end_str = part.split("-", 1)
            try:
                start = int(start_str)
                end = int(end_str)
            except ValueError as exc:
                raise ValueError(f"Invalid page range: {part}") from exc
            if start <= 0 or end <= 0 or end < start:
                raise ValueError(f"Invalid page range: {part}")
            for page in range(start, end + 1):
                pages.add(page - 1)
        else:
            try:
                page = int(part)
            except ValueError as exc:
                raise ValueError(f"Invalid page number: {part}") from exc
            if page <= 0:
                raise ValueError(f"Invalid page number: {part}")
            pages.add(page - 1)
    return sorted(pages)


def _expand_inputs(input_files: list[str]) -> list[str]:
    expanded: list[str] = []
    for item in input_files:
        matches = glob.glob(item)
        if matches:
            expanded.extend(matches)
        else:
            expanded.append(item)
    return expanded


def _normalize_headers(values: list[Any]) -> list[str]:
    headers: list[str] = []
    used: set[str] = set()
    for idx, raw in enumerate(values):
        name = "" if raw is None else str(raw).strip()
        if not name:
            name = f"column_{idx + 1}"
        base = name
        counter = 2
        while name in used:
            name = f"{base}_{counter}"
            counter += 1
        used.add(name)
        headers.append(name)
    return headers


def _table_to_dataframe(table: list[list[Any]]) -> pd.DataFrame:
    if not table:
        return pd.DataFrame()
    header = _normalize_headers(list(table[0]))
    rows = table[1:] if len(table) > 1 else []
    if not rows:
        return pd.DataFrame(columns=header)
    return pd.DataFrame(rows, columns=header)


def _text_to_dataframe(lines: list[str]) -> pd.DataFrame:
    cleaned = [line for line in (line.strip() for line in lines) if line]
    return pd.DataFrame({"text": cleaned})


def _write_dataframe(df: pd.DataFrame, output_format: str, target: str) -> None:
    if output_format == "csv":
        df.to_csv(target, index=False, encoding="utf8")
        return
    if output_format == "json":
        with open(target, "w", encoding="utf8") as handle:
            handle.write(df.to_json(orient="records", force_ascii=False, indent=2))
            handle.write("\n")
        return
    if output_format == "ndjson":
        df.to_json(target, orient="records", lines=True, force_ascii=False)
        return
    if output_format == "parquet":
        df.to_parquet(target, index=False)
        return
    raise ValueError(f"Unsupported output format: {output_format}")


def _write_dataframe_stdout(df: pd.DataFrame, output_format: str) -> None:
    if output_format == "csv":
        df.to_csv(sys.stdout, index=False)
        return
    if output_format == "json":
        sys.stdout.write(df.to_json(orient="records", force_ascii=False, indent=2))
        sys.stdout.write("\n")
        return
    if output_format == "ndjson":
        sys.stdout.write(df.to_json(orient="records", lines=True, force_ascii=False))
        sys.stdout.write("\n")
        return
    raise ValueError(f"Output format {output_format} requires --output or --output-dir")


class Extractor:
    """Document extraction command handler."""

    def extract(self, input_files: list[str], options: Optional[dict[str, Any]] = None) -> None:
        if options is None:
            options = {}
        if not input_files:
            raise ValueError("No input files provided.")

        output_format = _normalize_output_format(get_option(options, "output_format"))
        if output_format not in SUPPORTED_OUTPUTS:
            raise ValueError(f"Unsupported output format: {output_format}")

        output = get_option(options, "output")
        output_dir = get_option(options, "output_dir")
        method = _normalize_method(get_option(options, "method"))
        pages = _parse_pages(get_option(options, "pages"))
        ocr = bool(get_option(options, "ocr"))
        flatten = bool(get_option(options, "flatten"))

        if ocr:
            method = "ocr"

        inputs = _expand_inputs(input_files)
        if not inputs:
            raise ValueError("No input files matched the provided patterns.")

        from ..common.errors import (
            FileNotFoundError,
            FormatError,
            PermissionError,
            find_similar_files,
        )
        from ..common.path_utils import validate_file_path

        tables: list[dict[str, Any]] = []
        for input_path in inputs:
            # Validate file exists and is readable
            try:
                validate_file_path(input_path, check_read=True)
            except FileNotFoundError as e:
                suggestions = find_similar_files(input_path)
                raise FileNotFoundError(input_path, suggestions) from e
            except PermissionError as e:
                raise PermissionError(input_path, operation="read") from e

            filetype = os.path.splitext(input_path)[-1].lower().lstrip(".")
            if filetype not in SUPPORTED_INPUTS:
                raise FormatError(input_path, filetype, list(SUPPORTED_INPUTS.keys()))
            tables.extend(self._extract_file(input_path, filetype, method, pages, flatten))

        if output_format == "datapackage":
            self._write_datapackage(tables, output, output_dir)
            return

        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            self._write_tables_to_dir(tables, output_format, output_dir)
            return

        if output:
            if len(tables) != 1:
                raise ValueError("Multiple tables extracted; use --output-dir instead of --output.")
            _write_dataframe(tables[0]["df"], output_format, output)
            return

        if len(tables) != 1:
            raise ValueError("Multiple tables extracted; use --output-dir to write files.")
        _write_dataframe_stdout(tables[0]["df"], output_format)

    def _extract_file(
        self,
        input_path: str,
        filetype: str,
        method: Optional[str],
        pages: Optional[list[int]],
        flatten: bool,
    ) -> list[dict[str, Any]]:
        if filetype == "pdf":
            tables = self._extract_pdf(input_path, method, pages)
        elif filetype in {"doc", "docx"}:
            tables = self._extract_doc(input_path, filetype, method)
        else:
            tables = self._extract_excel(input_path)

        if flatten and tables:
            combined = []
            for idx, table in enumerate(tables, start=1):
                df = table["df"].copy()
                df["source_table_index"] = idx
                df["source_page"] = table.get("page")
                combined.append(df)
            merged = pd.concat(combined, ignore_index=True) if combined else pd.DataFrame()
            return [
                {
                    "df": merged,
                    "source": table.get("source", input_path),
                    "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_flattened",
                }
            ]

        return tables

    def _extract_pdf(
        self,
        input_path: str,
        method: Optional[str],
        pages: Optional[list[int]],
    ) -> list[dict[str, Any]]:
        try:
            import pdfplumber  # type: ignore
        except Exception as exc:
            raise ImportError(
                'PDF extraction requires pdfplumber. Install with `pip install "undatum[extract]"`.'
            ) from exc

        if method is None:
            method = "tables"
        if method not in {"tables", "text", "ocr"}:
            raise ValueError(f"Unsupported PDF method: {method}")

        tables: list[dict[str, Any]] = []
        with pdfplumber.open(input_path) as pdf:
            page_indices = pages if pages is not None else list(range(len(pdf.pages)))
            if method == "ocr":
                return self._extract_pdf_ocr(input_path, page_indices)
            for page_index in page_indices:
                if page_index >= len(pdf.pages):
                    continue
                page = pdf.pages[page_index]
                if method == "text":
                    text = page.extract_text() or ""
                    df = _text_to_dataframe(text.splitlines())
                    tables.append(
                        {
                            "df": df,
                            "page": page_index + 1,
                            "source": input_path,
                            "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_p{page_index + 1}_text",
                        }
                    )
                    continue
                extracted = page.extract_tables() or []
                for idx, table in enumerate(extracted, start=1):
                    df = _table_to_dataframe(table)
                    tables.append(
                        {
                            "df": df,
                            "page": page_index + 1,
                            "source": input_path,
                            "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_p{page_index + 1}_t{idx}",
                        }
                    )

        if method == "tables" and not tables:
            raise ValueError("No tables detected; try --method text or --method ocr.")
        return tables

    def _extract_pdf_ocr(self, input_path: str, page_indices: list[int]) -> list[dict[str, Any]]:
        try:
            import pytesseract  # type: ignore
            from pdf2image import convert_from_path  # type: ignore
        except Exception as exc:
            raise ImportError(
                "OCR extraction requires pdf2image and pytesseract. "
                'Install with `pip install "undatum[extract]"`.'
            ) from exc

        if not page_indices:
            return []

        tables: list[dict[str, Any]] = []
        for page_index in page_indices:
            images = convert_from_path(
                input_path, first_page=page_index + 1, last_page=page_index + 1
            )
            text_chunks = []
            for image in images:
                text_chunks.append(pytesseract.image_to_string(image))
            lines = "\n".join(text_chunks).splitlines()
            df = _text_to_dataframe(lines)
            tables.append(
                {
                    "df": df,
                    "page": page_index + 1,
                    "source": input_path,
                    "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_p{page_index + 1}_ocr",
                }
            )
        return tables

    def _extract_doc(
        self,
        input_path: str,
        filetype: str,
        method: Optional[str],
    ) -> list[dict[str, Any]]:
        if method is None:
            method = "text" if filetype == "docx" else "text"
        if method not in {"tables", "text"}:
            raise ValueError(f"Unsupported DOC method: {method}")

        if filetype == "doc":
            return self._extract_doc_binary(input_path, method)

        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise ImportError(
                'DOCX extraction requires python-docx. Install with `pip install "undatum[extract]"`.'
            ) from exc

        document = Document(input_path)
        tables: list[dict[str, Any]] = []
        if method == "text":
            lines = [p.text for p in document.paragraphs if p.text.strip()]
            df = _text_to_dataframe(lines)
            tables.append(
                {
                    "df": df,
                    "source": input_path,
                    "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_text",
                }
            )
            return tables

        for idx, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            df = _table_to_dataframe(rows)
            tables.append(
                {
                    "df": df,
                    "source": input_path,
                    "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_t{idx}",
                }
            )

        if method == "tables" and not tables:
            raise ValueError("No tables detected; try --method text.")
        return tables

    def _extract_doc_binary(self, input_path: str, method: str) -> list[dict[str, Any]]:
        try:
            import textract  # type: ignore
        except Exception as exc:
            raise ImportError(
                "DOC extraction requires textract or convert to DOCX first. "
                'Install with `pip install "undatum[extract]"`.'
            ) from exc

        raw = textract.process(input_path)
        text = raw.decode("utf8", errors="replace")
        df = _text_to_dataframe(text.splitlines())
        return [
            {
                "df": df,
                "source": input_path,
                "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_text",
            }
        ]

    def _extract_excel(self, input_path: str) -> list[dict[str, Any]]:
        try:
            data = pd.read_excel(input_path, sheet_name=None)
        except Exception as exc:
            raise ValueError(f"Failed to read spreadsheet: {exc}") from exc

        tables: list[dict[str, Any]] = []
        for idx, (sheet, df) in enumerate(data.items(), start=1):
            tables.append(
                {
                    "df": df,
                    "source": input_path,
                    "name": f"{os.path.splitext(os.path.basename(input_path))[0]}_{sheet or idx}",
                }
            )
        return tables

    def _write_tables_to_dir(
        self, tables: list[dict[str, Any]], output_format: str, output_dir: str
    ) -> None:
        for idx, table in enumerate(tables, start=1):
            name = table.get("name") or f"table_{idx}"
            filename = f"{name}{OUTPUT_EXT[output_format]}"
            target = os.path.join(output_dir, filename)
            _write_dataframe(table["df"], output_format, target)

    def _write_datapackage(
        self,
        tables: list[dict[str, Any]],
        output: Optional[str],
        output_dir: Optional[str],
    ) -> None:
        if not output and not output_dir:
            raise ValueError("datapackage output requires --output or --output-dir")

        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        output_file = output
        if output_file and output_dir:
            output_file = os.path.join(output_dir, os.path.basename(output_file))
        if not output_file:
            output_file = os.path.join(output_dir or os.getcwd(), "datapackage.json")

        with tempfile.TemporaryDirectory() as temp_dir:
            resources_dir = output_dir or temp_dir
            os.makedirs(resources_dir, exist_ok=True)
            resource_files: list[str] = []
            for idx, table in enumerate(tables, start=1):
                name = table.get("name") or f"resource_{idx}"
                filename = f"{name}.csv"
                target = os.path.join(resources_dir, filename)
                _write_dataframe(table["df"], "csv", target)
                resource_files.append(target)

            Packager().create(resource_files, {"output": output_file, "package_dir": output_dir})
